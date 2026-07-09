import streamlit as st
import docx
import re

# 網頁標題設定
st.set_page_config(page_title="慈文國中數資班專題課 word測驗閱卷系統", page_icon="📝")
st.title("📝 慈文國中數資班專題課 word測驗閱卷系統")
st.subheader("請上傳您的作業檔案 (.docx)，系統將進行格式與文字內容審查。")

uploaded_file = st.file_uploader("選擇上傳 Word 檔案 (.docx)", type=["docx"])

# 定義標準的文字內容範本（用來比對學生的內文，排除標題旁的姓名日期）
STANDARD_PARAGRAPHS = [
    "一、目的:探討在純水、自來水及其他溶液內，草履蟲是否會表現不同的移動速率。",
    "二、器材:",
    "三、步驟:",
    "(一)以燒杯盛裝純水、自來水、運動飲料和稀釋紅墨水各100mL。",
    "(二)各滴入1 mL的草履蟲水樣，並均勻混合。",
    "(三)以滴管取含有草履蟲的純水，置於複式顯微鏡下觀察並錄影。",
    "(四)以Tracker分析五隻草履蟲的移動速率，並記錄之。",
    "(五)將純水依序換成自來水、運動飲料、稀釋紅墨水，重複步驟(二)~(四)。"
]

# 表格內標準的八個器材文字
STANDARD_TABLE_CELLS = [
    "(一)含有草履蟲的水樣 1杯", "(二)滴管 1支",
    "(三)顯微鏡(含錄影裝備) 1組", "(四)純水 100mL",
    "(五)自來水 100mL", "(六)運動飲料 100mL",
    "(七)1:100稀釋紅墨水 100mL", "(八)Tracker(分析運動速率軟體) 1個"
]

def format_check(file_path, filename):
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        return None, [], [f"檔案解析失敗，請確保是標準的 docx 檔案。錯誤訊息: {e}"]

    score = 100
    passed_items = []
    failed_items = []

    # 1. 檢查檔名：必須符合 7X-XXX-文件三寶測驗 或 7X-XX-文件三寶測驗 (X限英文字與繁體中文，不含數字)
    pure_filename = filename.rsplit('.', 1)[0]
    filename_pattern = r"^7[a-zA-Z\u4e00-\u9fa5]-[a-zA-Z\u4e00-\u9fa5]+-文件三寶測驗$"
    
    if re.match(filename_pattern, pure_filename):
        passed_items.append("檔名規範：符合格式要求。")
    else:
        reason = "格式不符"
        if "文書三寶" in pure_filename:
            reason = "把「文件三寶」打成「文書三寶」了"
        elif any(char.isdigit() for char in pure_filename if char != '7'):
            reason = "檔名除了開頭的 7 以外，其餘位置包含了不允許的「數字」"
        failed_items.append(f"檔名規範：不符合格式（錯誤原因：{reason}）。正確應如 7E-解答-文件三寶測驗。 (-5分)")
        score -= 5

    # 2. 檢查基本資訊（標題與姓名日期）
    title_p = None
    student_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    for p in doc.paragraphs:
        if "不同溶液" in p.text:
            title_p = p
            break
            
    if title_p:
        text = title_p.text
        cleaned_text = text.replace("不同溶液中草履蟲的移動速度", "").replace("不同溶液下，草履蟲的移動速度", "").strip()
        cleaned_text = cleaned_text.replace(":", "").replace("：", "").strip()
        
        if len(cleaned_text) >= 2:
            passed_items.append("基本資訊：標題旁已正確填寫實驗者姓名與日期資訊。")
        else:
            failed_items.append(f"基本資訊：標題列「{text}」旁邊未包含或未完整寫出「實驗者姓名」與「今天日期」。 (-5分)")
            score -= 5
    else:
        failed_items.append("基本資訊：找不到實驗標題「不同溶液中草履蟲的移動速度」。 (-5分)")
        score -= 5

    # 3. 檢查全文文字內容、標點、換行
    content_errors = []
    filtered_student_paras = [p for p in student_paragraphs if "不同溶液" not in p]
    
    if len(filtered_student_paras) < len(STANDARD_PARAGRAPHS):
        content_errors.append(f"文字與換行：偵測到段落缺失或漏掉換行。標準應有 {len(STANDARD_PARAGRAPHS)} 個段落，但僅偵測到 {len(filtered_student_paras)} 個段落。")
    else:
        for idx, std_text in enumerate(STANDARD_PARAGRAPHS):
            if idx < len(filtered_student_paras):
                stud_text = filtered_student_paras[idx]
                if stud_text != std_text:
                    content_errors.append(f"段落內容不符 -> 正確應該是:「{std_text}」 | 您的檔案寫成:「{stud_text}」")

    # 4. 檢查表格內的器材文字
    table_cells_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in table_cells_text:
                    table_cells_text.append(cell_text)
                    
    for std_cell in STANDARD_TABLE_CELLS:
        if std_cell not in table_cells_text:
            content_errors.append(f"器材表格文字不符 -> 找不到或打錯了器材名稱:「{std_cell}」")

    if not content_errors:
        passed_items.append("文字與標點：全文文字、標點符號、換行完全與題目要求一致，無任何增減。")
    else:
        for err in content_errors[:3]:
            failed_items.append(f"文字內容錯誤：{err} (-5分)")
            score -= 5

    # 5. 檢查邊界與行距
    margin_ok = True
    if len(doc.sections) > 0:
        section = doc.sections[0]
        top = section.top_margin.cm if section.top_margin else 0
        bottom = section.bottom_margin.cm if section.bottom_margin else 0
        left = section.left_margin.cm if section.left_margin else 0
        right = section.right_margin.cm if section.right_margin else 0
        
        if not (2.45 <= top <= 2.55 and 2.45 <= bottom <= 2.55 and 2.45 <= left <= 2.55 and 2.45 <= right <= 2.55):
            margin_ok = False
            failed_items.append(f"版面邊界：上下左右邊界未設定為 2.5cm（實際偵測值為 上:{round(top,2)}cm, 下:{round(bottom,2)}cm, 左:{round(left,2)}cm, 右:{round(right,2)}cm）。 (-5分)")
            score -= 5

    spacing_ok = True
    for p in doc.paragraphs:
        if p.text.strip() and p.paragraph_format.line_spacing:
            if p.paragraph_format.line_spacing != 1.0 and p.paragraph_format.line_spacing_rule is not None:
                spacing_ok = False
                break
                
    if margin_ok and spacing_ok:
        passed_items.append("版面邊界與行距：邊界 2.5cm 且行距為 1.0 行設定正確。")
    elif margin_ok and not spacing_ok:
        failed_items.append("行距設定：部分內文段落行距非 1.0 行。 (-5分)")
        score -= 5

    # 6. 檢查字體與顏色
    font_ok = True
    color_ok = True
    wrong_fonts = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for run in p.runs:
            if not run.text.strip():
                continue
            if run.font.name and "標楷體" not in run.font.name and "KaiTi" not in run.font.name and "Times New Roman" not in run.font.name:
                font_ok = False
                if run.text.strip()[:10] not in wrong_fonts:
                    wrong_fonts.append(f"「{run.text.strip()[:6]}」")
            if run.font.color and run.font.color.rgb and run.font.color.rgb != docx.shared.RGBColor(0,0,0):
                color_ok = False

    if font_ok:
        passed_items.append("字體規範：中文字為標楷體，英文與數字皆為 Times New Roman。")
    else:
        failed_items.append(f"字體規範：發現部分文字字體設定錯誤（例如：{', '.join(wrong_fonts[:3])} 未正確設定）。 (-5分)")
        score -= 5

    if color_ok:
        passed_items.append("文字顏色：所有文字與數字顏色皆為黑色。")
    else:
        failed_items.append("文字顏色：發現部分文字顏色非黑色. (-5分)")
        score -= 5

    # 7. 檢查字級大小
    size_ok = True
    wrong_sizes = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        is_title_p = "不同溶液" in p.text
        for run in p.runs:
            if not run.text.strip() or run.font.size is None:
                continue
            size = run.font.size.pt
            if is_title_p and size != 14.0:
                size_ok = False
                wrong_sizes.append(f"標題「{run.text[:6]}」(偵測為 {size}號)")
            if not is_title_p and size != 12.0:
                size_ok = False
                if f"內文「{p.text[:6]}...」" not in wrong_sizes:
                    wrong_sizes.append(f"內文「{p.text[:6]}...」(偵測為 {size}號)")

    if size_ok:
        passed_items.append("字級大小：標題為 14 號字，其餘內文皆為 12 號字。")
    else:
        failed_items.append(f"字級大小：發現字號設定錯誤（例如：{', '.join(wrong_sizes[:2])}，規定標題14號、內文12號）。 (-5分)")
        score -= 5

    # 8. 檢查粗體限制
    bold_ok = True
    bold_targets = ["不同溶液", "一、目的", "二、器材", "三、步驟"]
    missing_bold = []
    extra_bold = []
    
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        is_target = any(target in p.text for target in bold_targets)
        has_any_bold_run = any(run.bold is True for run in p.runs if run.text.strip() and run.text.strip() not in [":", "："])
        is_style_bold = "Heading" in p.style.name or "Heading" in str(p.style.base_style)
        paragraph_is_bold = has_any_bold_run or is_style_bold
        
        if is_target and not paragraph_is_bold:
            bold_ok = False
            if f"「{p.text[:6]}...」" not in missing_bold:
                missing_bold.append(f"「{p.text[:6]}...」")
        elif not is_target and paragraph_is_bold:
            if all(run.bold is True for run in p.runs if run.text.strip()):
                bold_ok = False
                if f"「{p.text[:6]}...」" not in extra_bold:
                    extra_bold.append(f"「{p.text[:6]}...」")

    if bold_ok:
        passed_items.append("粗體限制：只有指定的標題部分為粗體，其餘內文皆為正常體。")
    else:
        err_msg = "粗體限制：粗體設定不符規範。"
        if missing_bold:
            err_msg += f" 應設為粗體卻漏設的有：{', '.join(missing_bold)}；"
        if extra_bold:
            err_msg += f" 不應設為粗體卻誤設的有：{', '.join(extra_bold)}；"
        failed_items.append(f"{err_msg} (-5分)")
        score -= 5

    # 9. 檢查器材呈現（2列4欄表格與框線）
    if len(doc.tables) > 0:
        table = doc.tables[0]
        rows = len(table.rows)
        cols = len(table.columns)
        table_shape_ok = (rows == 2 and cols == 4)
        table_style_ok = "Border" in table.style.name or "Normal" in table.style.name or "Table Grid" in table.style.name
        
        if table_shape_ok:
            passed_items.append(f"器材表格：已使用 2×4 表格呈現八個器材。")
        else:
            failed_items.append(f"器材表格：表格行列數不對（規定應為 2列×4欄，系統在您的檔案偵測到為 {rows}列×{cols}欄）。 (-5分)")
            score -= 5
            
        if table_style_ok:
            passed_items.append("表格框線：表格已設定為無框線/隱藏框線。")
        else:
            failed_items.append("表格框線：未將表格調整為「無框線」。 (-5分)")
            score -= 5
    else:
        failed_items.append("器材表格：文件中完全找不到器材表格（應以 2列×4欄 表格呈現）。 (-5分)")
        score -= 10

    # 10. 檢查步驟縮排 (已完美整合成單一連續區塊)
    indent_ok = True
    step_patterns = ["(一)", "(二)", "(三)", "(四)", "(五)", "（一）", "（二）", "（三）", "（四）", "（五）"]
    missing_indent_steps = []
    
    for p in doc.paragraphs:
        matched_pattern = next((pat for pat in step_patterns if p.text.strip().startswith(pat)), None)
        if matched_pattern:
            # 判斷有無設定左側縮排或首行縮排
            has_indent = (p.paragraph_format.left_indent and p.paragraph_format.left_indent.pt > 0) or \
                         (p.paragraph_format.first_line_indent and p.paragraph_format.first_line_indent.pt > 0)
            if has_indent:
                pass
            else:
                indent_ok = False
                missing_indent_steps.append(matched_pattern)

        if indent_ok and not missing_indent_steps:
            passed_items.append("步驟縮排：三、步驟底下的項目已正確設定縮排。")
        else:
            failed_items.append(f"步驟縮排：三、步驟底下的項目未設定「縮排 2 字元」（系統偵測到未縮排的步驟有：{', '.join(missing_indent_steps) if missing_indent_steps else '所有步驟'}）。 (-5分)")
            score -= 5

    score = max(0, score)
    return score, passed_items, failed_items
    if uploaded_file is not None:
	filename = uploaded_file.namest.info(f"📁 已偵測到上傳檔名：{filename}")
    with st.spinner("系統正在進行格式與內文審查..."):
        result = format_check(uploaded_file, filename)
    if result is None:st.error("檔案解析失敗")
    else:
	score, passed, failed = result
	st.write("---")
	st.markdown("### 👨‍🏫 閱卷老師評分報告")
    if score >= 90:
 	st.success(f"### 最終得分：{score} 分")
    elif score >= 60:
	st.warning(f"### 最終得分：{score} 分")
    else:
	st.error(f"### 最終得分：{score} 分")
    st.markdown("#### 🟢 合格項目")
    if passed:
	for item in passed:st.write(f" - {item}")
    else:st.write("無項目合格。")
    st.markdown("#### 🔴 不合格項目與扣分明細")
    if failed:
	for item in failed:
		st.write(f" - {item}")
    else:st.write("🎉 太棒了！完美符合所有格式要求與文字內容，沒有任何扣分！")